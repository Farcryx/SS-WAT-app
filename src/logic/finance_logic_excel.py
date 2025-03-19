import os
import sys

import openpyxl
from pypdf import PdfReader
import datetime
import re
import msoffcrypto
import io
from docx import Document
import customtkinter
import tkinter

# Define values for reading excel file
min_row = 1
max_row = 10
min_col = 1
max_col = 20
sheet_names = None


def read_titles(workbook, active_sheet_number):
    column_titles = []
    global sheet_names
    active_sheet = workbook[sheet_names[active_sheet_number]]
    for col in range(min_col, max_col):
        if active_sheet.cell(1, column=col).value is not None:
            column_titles.append([col, active_sheet.cell(1, column=col).value])
            print(active_sheet.cell(1, column=col).value)
    return column_titles


def read_column(workbook, active_sheet_number, column_nr):
    list_of_cells_in_column = []
    global sheet_names
    active_sheet = workbook[sheet_names[active_sheet_number]]
    row = active_sheet.max_row
    for row in range(1 + 1, row):
        if active_sheet.cell(row=row, column=column_nr).value:
            list_of_cells_in_column.append([row, active_sheet.cell(row=row, column=column_nr).value])
    return list_of_cells_in_column


def read_excel(file_path):
    try:
        workbook = openpyxl.load_workbook(file_path)
    except FileNotFoundError:
        print(f"Plik {file_path} nie został znaleziony.")
        return

    global sheet_names
    sheet_names = workbook.sheetnames
    return sheet_names, workbook

def read_pdf(file_path):
    reader = PdfReader(file_path)
    print(len(reader.pages))
    pdf_content = ""
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pdf_content += text
    # Use regex to find all standalone five-digit numbers
    list_of_albums_from_pdf = re.findall(r'\b\d{5}\b', pdf_content)
    # make data in list unique
    list_of_albums_from_pdf = list(set(list_of_albums_from_pdf))
    print(f"Lista albumów z pdf: {list_of_albums_from_pdf}")
    return list_of_albums_from_pdf


def read_docx(file_path, password, log_file):
    try:
        # First try to open normally
        try:
            document = Document(file_path)
        except:
            # If fails, file might be encrypted - try with password
            with open(file_path, "rb") as f:
                office_file = msoffcrypto.OfficeFile(f)
                office_file.load_key(password=password)
                decrypted = io.BytesIO()
                office_file.decrypt(decrypted)
                document = Document(decrypted)
            
        # Now read the first table with two columns: 
        # in first cell -> kwota (amount)
        # in second cell -> dane uczestnika; numer albumu jest ostatnim słowem
        album_dict = {}
        for table in document.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    amount = row.cells[0].text.strip()
                    data = row.cells[1].text.strip()
                    if data:
                        # Album number is the last word in the cell
                        album_number = data.split()[-1]
                        if len(album_number) == 5 and amount != '':
                            album_dict[album_number] = amount[1:]
                        elif len(album_number) == 5 and amount == '':
                            print(f"Nie znaleziono kwoty dla osoby: {data.split()[0]} {data.split()[1]}")
                            write_to_log_file(log_file, f"{datetime.datetime.now()} - Nie znaleziono kwoty dla osoby: {data.split()[0]} {data.split()[1]}\n")
                            album_dict[album_number] = "Zapłacone"
                        else:
                            print(f"Nie znaleziono numeru albumu dla osoby: {data.split()[0]} {data.split()[1]}")
                            write_to_log_file(log_file, f"{datetime.datetime.now()} - Nie znaleziono numeru albumu dla osoby: {data.split()[0]} {data.split()[1]}\n")

        print(f"Znaleziono {len(album_dict)} albumów: {album_dict}")
        # check if number of rows of document table is equal to lenght of dict
        if len(document.tables[0].rows) == len(album_dict):
            print(f"Zrealizowano pomyślnie całą weryfikację!")
            write_to_log_file(log_file, f"{datetime.datetime.now()} - Zrealizowano pomyślnie całą weryfikację!\n")
        else:
            print(f"Liczba płatności {len(document.tables[0].rows) - 1} nie zgadza się z liczbą zweryfikowanych {len(album_dict)}")
            write_to_log_file(log_file, f"{datetime.datetime.now()} - Liczba płatności {len(document.tables[0].rows) - 1} nie zgadza się z liczbą zweryfikowanych {len(album_dict)}\n")
        return album_dict
    except Exception as e:
        print(f"Error reading docx file: {e}")
        return {}


def read_document(file_path, password, log_file):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_docx(file_path, password, log_file)
    else:
        print("Unsupported file format!")
        return []

def write_to_log_file(log_file, message=None):
    if message and log_file:
        try:
            log_file.write(message)
        except Exception as e:
            print(f"Nie udało się zapisać do pliku logów: {e}")
    else:
        print("Nie można zapisać pustej wiadomości lub log_file jest None.")
        log_file.close()

def create_log_file():
    # Create logs directory if it doesn't exist
    if not os.path.exists('logs'):
        try:
            os.makedirs('logs')
        except Exception as e:
            print(f"Nie udało się utworzyć katalogu logów: {e}")
            return None

    # Create log file with current date
    try:
        date = datetime.datetime.now().strftime("%Y-%m-%d %H-%M-%S")
        log_file = open(f"logs/log_{date}.log", "w")
        # Write basic information about the application and system
        log_file.write("Logi z aplikacji SS WAT z modułu płatności\n")
        log_file.write("Autor aplikacji: Łukasz Sokół\n")
        log_file.write(f"Data utworzenia logów: {date}\n")
        log_file.write(f"System operacyjny: {os.name}\n")
        log_file.write(f"Wersja Pythona: {os.sys.version}\n")
        log_file.write(f"Wersja openpyxl: {openpyxl.__version__}\n")
        # PdfReader, msoffcrypto, and Document do not have __version__ attributes; write 'unknown'
        log_file.write("Wersja pypdf: unknown\n")
        log_file.write("Wersja msoffcrypto: unknown\n")
        log_file.write("Wersja python-docx: unknown\n")
        log_file.write(f"Wersja customtkinter: {customtkinter.__version__}\n")
        log_file.write(f"Wersja Tkinter: {tkinter.TkVersion}\n")
        log_file.write(f"Ścieżka do pliku logów: {os.path.abspath(log_file.name)}\n")
        return log_file
    except Exception as e:
        print(f"Nie udało się utworzyć pliku logów: {e}")
        return None



def analyze_data(workbook, excel_path, document_path, number_of_active_sheet, index_album_column, index_payment_column,
                 log_file, doc_password=None):
    print(f"Analizuję dane z pliku {document_path}")
    write_to_log_file(log_file, f"{datetime.datetime.now()} - Analizuję dane z pliku {document_path}\n")
    # Get list/dictionary of albums from the document (pdf returns list; docx now returns dict)
    if doc_password is None:
        try:
            with open(".password", "r") as password_file:
                doc_password = password_file.read().strip()
        except FileNotFoundError:
            print("Nie znaleziono pliku z hasłem do dokumentu docx")
            write_to_log_file(log_file, f"{datetime.datetime.now()} - Nie znaleziono pliku z hasłem do dokumentu docx\n")
            return
    albums_result = read_document(document_path, doc_password, log_file)
    if isinstance(albums_result, dict):
        album_items = albums_result.items()
        print(f"Znaleziona liczba albumów w dokumencie {len(albums_result)}: {albums_result}")
        write_to_log_file(log_file,
                          f"{datetime.datetime.now()} - Znaleziona liczba albumów w dokumencie {len(albums_result)}\n")
        write_to_log_file(log_file, f"{datetime.datetime.now()} - Znalezione albumy: {albums_result}\n")
    else:
        # In case of PDF we still get a list
        album_items = [(album, "TAK") for album in albums_result]
        print(f"Znaleziona liczba albumów w dokumencie {len(albums_result)}: {albums_result}")
        write_to_log_file(log_file,
                          f"{datetime.datetime.now()} - Znaleziona liczba albumów w dokumencie {len(albums_result)}\n")
        write_to_log_file(log_file, f"{datetime.datetime.now()} - Znalezione albumy: {albums_result}\n")

    # Get column with album numbers from Excel
    album_column_from_excel = read_column(workbook, number_of_active_sheet, index_album_column)
    write_to_log_file(log_file,
                      f"{datetime.datetime.now()} - Znaleziona liczba albumów w pliku Excel {len(album_column_from_excel)}\n")
    write_to_log_file(log_file,
                      f"{datetime.datetime.now()} - Znalezione albumy w pliku Excel: {album_column_from_excel}\n")

    # Find albums from document in Excel and write the paid amount (or "TAK" for PDFs)
    found_albums = []
    not_found_albums = []
    for album, paid_value in album_items:
        found = False
        for row in album_column_from_excel:
            print(f"Porównuję album {album} z {row[1]}")
            if str(album) in str(row[1]):
                found_albums.append(album)
                workbook[sheet_names[number_of_active_sheet]].cell(row=row[0],
                                                                    column=index_payment_column).value = paid_value
                print(f"Wpisano {paid_value} w komórkę {row[0]} w kolumnie {index_payment_column}")
                write_to_log_file(log_file,
                                  f"{datetime.datetime.now()} - Wpisano {paid_value} w komórkę {row[0]} w kolumnie {index_payment_column}\n")
                found = True
                break
        if not found:
            not_found_albums.append(album)
            write_to_log_file(log_file, f"{datetime.datetime.now()} - Nie znaleziono albumu {album} w pliku Excel\n")

    workbook.save(excel_path)
    print(f"Pomyślnie zapisano zmiany w pliku Excel {excel_path}")
    write_to_log_file(log_file, f"{datetime.datetime.now()} - Pomyślnie zapisano zmiany w pliku Excel {excel_path}\n")
    write_to_log_file(log_file, f"Podsumowanie\n")
    write_to_log_file(log_file, f"Znaleziono {len(found_albums)} albumów w pliku Excel\n")
    write_to_log_file(log_file, f"Nie znaleziono {len(not_found_albums)} albumów w pliku Excel\n")
    write_to_log_file(log_file, f"Nieznalezione albumy: {not_found_albums}\n")
    # log_file.close()